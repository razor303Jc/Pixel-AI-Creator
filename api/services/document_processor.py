"""
Document Processing Service

Handles text extraction from various file formats (PDF, DOCX, TXT)
and provides text analysis capabilities including summarization and keyword extraction.
"""

import logging
import aiofiles
from pathlib import Path
from typing import Optional, List, Dict, Any
import asyncio
import io

# Document processing libraries
try:
    import PyPDF2
    import docx
    from docx import Document
except ImportError:
    PyPDF2 = None
    docx = None
    Document = None

# AI services for processing
from services.openai_service import OpenAIService

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Document processing service for text extraction and analysis."""

    def __init__(self):
        """Initialize the document processor."""
        self.openai_service = OpenAIService()

        # Check for required libraries
        if not PyPDF2:
            logger.warning("PyPDF2 not installed - PDF processing disabled")
        if not docx:
            logger.warning("python-docx not installed - DOCX processing disabled")

    async def extract_text(self, file_path: Path, file_extension: str) -> str:
        """
        Extract text content from various file formats.

        Args:
            file_path: Path to the file
            file_extension: File extension (.pdf, .docx, .txt)

        Returns:
            Extracted text content

        Raises:
            ValueError: If file format is not supported
            Exception: If text extraction fails
        """
        try:
            if file_extension == ".txt":
                return await self._extract_text_from_txt(file_path)
            elif file_extension == ".pdf":
                return await self._extract_text_from_pdf(file_path)
            elif file_extension in [".docx", ".doc"]:
                return await self._extract_text_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")

        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {str(e)}")
            raise

    async def _extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        try:
            async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                content = await f.read()
            return content.strip()
        except UnicodeDecodeError:
            # Try with different encoding
            async with aiofiles.open(file_path, "r", encoding="latin-1") as f:
                content = await f.read()
            return content.strip()

    async def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        if not PyPDF2:
            raise ImportError("PyPDF2 not installed - cannot process PDF files")

        def extract_pdf_sync():
            """Synchronous PDF extraction to run in thread pool."""
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""

                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"

                return text.strip()

        # Run in thread pool to avoid blocking
        loop = asyncio.get_event_loop()
        text = await loop.run_in_executor(None, extract_pdf_sync)

        if not text:
            raise Exception("No text could be extracted from PDF")

        return text

    async def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        if not docx:
            raise ImportError("python-docx not installed - cannot process DOCX files")

        def extract_docx_sync():
            """Synchronous DOCX extraction to run in thread pool."""
            doc = Document(file_path)
            text = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)

            return "\n".join(text)

        # Run in thread pool to avoid blocking
        loop = asyncio.get_event_loop()
        text = await loop.run_in_executor(None, extract_docx_sync)

        if not text:
            raise Exception("No text could be extracted from DOCX")

        return text

    async def generate_summary(self, text: str, max_length: int = 200) -> str:
        """
        Generate an AI summary of the document content.

        Args:
            text: Document text content
            max_length: Maximum summary length in words

        Returns:
            AI-generated summary
        """
        try:
            if len(text.split()) < 50:
                # Document is too short for summarization
                return text[:max_length] + "..." if len(text) > max_length else text

            prompt = f"""
            Please provide a concise summary of the following document in {max_length} words or less:

            {text[:4000]}  # Limit input to avoid token limits

            Summary:
            """

            summary = await self.openai_service.generate_response(
                prompt=prompt, max_tokens=300, temperature=0.3
            )

            return summary.strip()

        except Exception as e:
            logger.error(f"Summary generation failed: {str(e)}")
            # Fallback to simple truncation
            words = text.split()
            if len(words) > max_length:
                return " ".join(words[:max_length]) + "..."
            return text

    async def extract_keywords(self, text: str, max_keywords: int = 10) -> List[str]:
        """
        Extract key topics and terms from the document.

        Args:
            text: Document text content
            max_keywords: Maximum number of keywords to extract

        Returns:
            List of extracted keywords
        """
        try:
            prompt = f"""
            Extract {max_keywords} key topics, terms, and concepts from this document.
            Return only the keywords, separated by commas:

            {text[:3000]}  # Limit input to avoid token limits

            Keywords:
            """

            response = await self.openai_service.generate_response(
                prompt=prompt, max_tokens=100, temperature=0.3
            )

            # Parse keywords from response
            keywords = [
                keyword.strip().lower()
                for keyword in response.split(",")
                if keyword.strip()
            ]

            return keywords[:max_keywords]

        except Exception as e:
            logger.error(f"Keyword extraction failed: {str(e)}")
            # Fallback to simple word frequency analysis
            return self._extract_keywords_fallback(text, max_keywords)

    def _extract_keywords_fallback(self, text: str, max_keywords: int) -> List[str]:
        """
        Fallback keyword extraction using simple word frequency.

        Args:
            text: Document text content
            max_keywords: Maximum number of keywords to extract

        Returns:
            List of extracted keywords
        """
        try:
            # Simple word frequency analysis
            words = text.lower().split()

            # Remove common stop words
            stop_words = {
                "the",
                "a",
                "an",
                "and",
                "or",
                "but",
                "in",
                "on",
                "at",
                "to",
                "for",
                "of",
                "with",
                "by",
                "is",
                "are",
                "was",
                "were",
                "be",
                "been",
                "being",
                "have",
                "has",
                "had",
                "do",
                "does",
                "did",
                "will",
                "would",
                "could",
                "should",
                "may",
                "might",
                "must",
                "can",
                "this",
                "that",
                "these",
                "those",
                "i",
                "you",
                "he",
                "she",
                "it",
                "we",
                "they",
                "me",
                "him",
                "her",
                "us",
                "them",
            }

            # Count word frequencies
            word_count = {}
            for word in words:
                # Clean word and filter
                clean_word = "".join(c for c in word if c.isalnum())
                if (
                    len(clean_word) > 3
                    and clean_word not in stop_words
                    and not clean_word.isdigit()
                ):
                    word_count[clean_word] = word_count.get(clean_word, 0) + 1

            # Sort by frequency and return top keywords
            sorted_words = sorted(word_count.items(), key=lambda x: x[1], reverse=True)

            return [word for word, count in sorted_words[:max_keywords]]

        except Exception as e:
            logger.error(f"Fallback keyword extraction failed: {str(e)}")
            return []

    async def chunk_text(
        self, text: str, chunk_size: int = 1000, overlap: int = 100
    ) -> List[str]:
        """
        Split text into overlapping chunks for vector storage.

        Args:
            text: Text to chunk
            chunk_size: Size of each chunk in characters
            overlap: Overlap between chunks in characters

        Returns:
            List of text chunks
        """
        try:
            if len(text) <= chunk_size:
                return [text]

            chunks = []
            start = 0

            while start < len(text):
                end = start + chunk_size

                # Try to break at sentence boundaries
                if end < len(text):
                    # Look for sentence ending within the last 200 characters
                    sentence_end = -1
                    for i in range(max(0, end - 200), end):
                        if text[i] in ".!?":
                            sentence_end = i + 1
                            break

                    if sentence_end > start:
                        end = sentence_end

                chunk = text[start:end].strip()
                if chunk:
                    chunks.append(chunk)

                start = end - overlap

                # Prevent infinite loop
                if start >= end:
                    start = end

            return chunks

        except Exception as e:
            logger.error(f"Text chunking failed: {str(e)}")
            return [text]  # Return original text as single chunk
